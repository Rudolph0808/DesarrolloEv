"""
Generador del reporte DOCX para la Entrega 2 del compilador Little Duck.

Lee los archivos de casos_prueba/ y sus salidas en casos_prueba/salidas/
y produce Reporte_Entrega2.docx.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).parent
CASOS_DIR = ROOT / "casos_prueba"
SALIDAS_DIR = CASOS_DIR / "salidas"


# ---------- Helpers de formato ----------

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_code_block(doc, text, font_size=9):
    """Inserta un bloque de código con fondo gris claro y fuente monoespaciada."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F2F2F2')
    cell.paragraphs[0].text = ''
    for line in text.splitlines():
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(font_size)
    # Quitar el primer párrafo vacío
    first = cell.paragraphs[0]
    if not first.text.strip():
        first._element.getparent().remove(first._element)
    return table


def add_inline_code(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return run


def add_paragraph_with_inline(doc, segments):
    """segments: lista de tuplas (text, is_code)."""
    para = doc.add_paragraph()
    for txt, is_code in segments:
        if is_code:
            add_inline_code(para, txt)
        else:
            para.add_run(txt)
    return para


def set_global_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# ---------- Cargar archivos de prueba y salidas ----------

def load_case(name):
    code = (CASOS_DIR / f"{name}.txt").read_text(encoding='utf-8')
    out = (SALIDAS_DIR / f"{name}.out").read_text(encoding='utf-8')
    return code, out


def extract_errors(output):
    return [l for l in output.splitlines() if '[Error' in l]


def extract_quads_section(output):
    """Devuelve solo la sección de cuádruplos del output (acortada si es larga)."""
    lines = output.splitlines()
    start, end = None, None
    for i, l in enumerate(lines):
        if 'REPRESENTACIÓN INTERMEDIA' in l:
            start = i
        if start is not None and 'TABLA DE SÍMBOLOS' in l:
            end = i
            break
    if start is None:
        return ""
    section = lines[start:end] if end else lines[start:]
    return "\n".join(section)


def extract_symbols_section(output):
    lines = output.splitlines()
    start = None
    for i, l in enumerate(lines):
        if 'TABLA DE SÍMBOLOS' in l:
            start = i
            break
    if start is None:
        return ""
    end = len(lines)
    for j in range(start, len(lines)):
        if '[OK] Representación' in lines[j]:
            end = j
            break
    return "\n".join(lines[start:end]).strip()


# ---------- Construcción del reporte ----------

def build_portada(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tecnológico de Monterrey")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph("Diseño de Compiladores  •  Semestre Ene–Jun 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Compilador Little Duck")
    run.bold = True
    run.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Entrega 2 — Análisis semántico y representación intermedia")
    run.font.size = Pt(14)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Integrantes:")
    run.bold = True

    for n in ("[Integrante 1]", "[Integrante 2]", "[Integrante 3]", "[Integrante 4]"):
        p = doc.add_paragraph(n)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph("Fecha de entrega: [Fecha]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def build_introduccion(doc):
    doc.add_heading("1. Introducción", level=1)

    doc.add_heading("1.1 Descripción del problema", level=2)
    doc.add_paragraph(
        "Little Duck es un lenguaje de programación sencillo, diseñado con fines didácticos "
        "para el curso de Diseño de Compiladores. Su gramática soporta declaración de "
        "variables tipadas (int, float, string), funciones con parámetros y retorno, "
        "estructuras de control (if/else, do-while con break), expresiones aritméticas y "
        "relacionales, asignaciones, impresión por consola y llamadas a funciones."
    )
    doc.add_paragraph(
        "El objetivo del proyecto es construir un compilador completo para Little Duck en "
        "Python utilizando PLY (Python Lex-Yacc). La entrega 2 extiende el trabajo previo "
        "(léxico + sintaxis) agregando el análisis semántico y la generación de la "
        "representación intermedia en forma de cuádruplos. Si el programa de entrada "
        "contiene errores léxicos, sintácticos o semánticos, el compilador los reporta con "
        "mensajes descriptivos; si no hay errores, produce los cuádruplos correspondientes "
        "tanto en pantalla como en un archivo de salida prueba-ir.txt."
    )

    doc.add_heading("1.2 Cambios respecto a la entrega 1", level=2)
    doc.add_paragraph(
        "La entrega 1 cubría únicamente la fase léxica y sintáctica: tokens, palabras "
        "reservadas, gramática LALR(1) y reporte de errores léxicos y sintácticos con "
        "estrategias de recuperación. La entrega 2 incorpora las siguientes funcionalidades "
        "nuevas:"
    )

    cambios = [
        "Estructuras de datos semánticas: cubo semántico, directorio de funciones, "
        "tablas de variables por scope, pilas de operandos, operadores, tipos y saltos.",
        "Soporte para la instrucción break dentro de ciclos.",
        "Cuatro validaciones semánticas: tipos, declaraciones y scopes, llamadas a "
        "funciones (con verificación de tipos y aridad), y control de flujo "
        "(condiciones booleanas y uso correcto de break).",
        "Generación de cuádruplos para expresiones, asignaciones, llamadas a funciones, "
        "control de flujo y operaciones de salida.",
        "Salida formateada en pantalla con cuádruplos numerados desde 1, columna del "
        "tipo del resultado, y tabla de símbolos completa.",
        "Producción del archivo prueba-ir.txt con la representación intermedia.",
        "Limpieza del lenguaje: eliminación de bool/true/false del set declarable, "
        "manteniendo bool únicamente como tipo interno de resultados de comparaciones.",
        "Lectura literal del archivo de entrada con input = open(\"prueba.txt\").read(), "
        "siguiendo la especificación de la entrega.",
    ]
    for c in cambios:
        doc.add_paragraph(c, style='List Bullet')


def build_metodologia(doc):
    doc.add_heading("2. Metodología", level=1)

    doc.add_heading("2.1 Arquitectura general", level=2)
    doc.add_paragraph(
        "El compilador está implementado en un único archivo main.py, organizado en cinco "
        "secciones:"
    )
    arq = [
        "Sección 0 — Estructuras de datos semánticas (clases y constantes de tipos).",
        "Sección 1 — Léxico: tokens, palabras reservadas y expresiones regulares.",
        "Sección 2 — Manejo de errores léxicos y sintácticos con reporte de línea y columna.",
        "Sección 3 — Gramática LR con acciones semánticas y neuralgic points.",
        "Sección 4 — Ejecución, formato de salida y escritura del archivo prueba-ir.txt.",
    ]
    for a in arq:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_paragraph(
        "El pipeline de compilación es el siguiente: el archivo prueba.txt se lee como "
        "texto plano; el lexer (ply.lex) produce un stream de tokens; el parser (ply.yacc) "
        "ejecuta una gramática LALR(1) y, en puntos clave de cada producción, dispara "
        "acciones semánticas sobre el SemanticManager. Estas acciones registran "
        "declaraciones, validan tipos a través del cubo semántico, manejan pilas de "
        "operandos/operadores/tipos/saltos, generan cuádruplos y aplican backpatching para "
        "los saltos pendientes. Al finalizar el parseo se imprime el resultado y se escribe "
        "el archivo de salida."
    )

    doc.add_heading("2.2 Estructuras de datos", level=2)

    # SemanticCube
    doc.add_heading("Cubo semántico (SemanticCube)", level=3)
    doc.add_paragraph(
        "Es un diccionario indexado por la terna (operador, tipo_izquierdo, tipo_derecho) "
        "que devuelve el tipo del resultado, o el valor 'error' si la operación no es "
        "válida para esos tipos. Implementa también la regla de asignación (=), donde "
        "asignar float a int produce error pero asignar int a float es legal."
    )
    add_code_block(doc, """class SemanticCube:
    def __init__(self):
        self.cube = {}
        arith = ['+', '-', '*', '/']
        for op in arith:
            self.cube[(op, T_INT, T_INT)] = T_INT
            self.cube[(op, T_FLOAT, T_FLOAT)] = T_FLOAT
            self.cube[(op, T_INT, T_FLOAT)] = T_FLOAT
            self.cube[(op, T_FLOAT, T_INT)] = T_FLOAT
        for op in ['==', '!=', '>', '<', '>=', '<=']:
            self.cube[(op, T_INT, T_INT)] = T_BOOL
            self.cube[(op, T_FLOAT, T_FLOAT)] = T_BOOL
            self.cube[(op, T_INT, T_FLOAT)] = T_BOOL
            self.cube[(op, T_FLOAT, T_INT)] = T_BOOL
        # Igualdad de strings
        for op in ['==', '!=']:
            self.cube[(op, T_STRING, T_STRING)] = T_BOOL
        # Asignación
        self.cube[('=', T_INT, T_INT)] = T_INT
        self.cube[('=', T_FLOAT, T_FLOAT)] = T_FLOAT
        self.cube[('=', T_FLOAT, T_INT)] = T_FLOAT
        self.cube[('=', T_STRING, T_STRING)] = T_STRING

    def query(self, op, t1, t2):
        return self.cube.get((op, t1, t2), T_ERROR)
""")

    doc.add_paragraph("Tabla resumida del cubo semántico:")
    cube_table = doc.add_table(rows=1, cols=4)
    cube_table.style = 'Light Grid Accent 1'
    hdr = cube_table.rows[0].cells
    hdr[0].text = "Operador"
    hdr[1].text = "Tipo Izq."
    hdr[2].text = "Tipo Der."
    hdr[3].text = "Resultado"

    cube_rows = [
        ("+, -, *, /", "int", "int", "int"),
        ("+, -, *, /", "float", "int", "float"),
        ("+, -, *, /", "int", "float", "float"),
        ("+, -, *, /", "float", "float", "float"),
        ("+, -, *, /", "string", "*", "error"),
        (">, <, >=, <=", "int", "float", "bool"),
        ("==, !=", "int", "int", "bool"),
        ("==, !=", "string", "string", "bool"),
        ("=", "int", "int", "int"),
        ("=", "float", "int", "float"),
        ("=", "int", "float", "error"),
        ("=", "string", "string", "string"),
    ]
    for row in cube_rows:
        cells = cube_table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    # FunctionDirectory / VariableTable
    doc.add_heading("Directorio de funciones y tabla de variables", level=3)
    doc.add_paragraph(
        "El directorio de funciones (FunctionDirectory) es un diccionario indexado por el "
        "nombre de la función. Cada entrada contiene el tipo de retorno, la lista ordenada "
        "de parámetros, una VariableTable local con las variables del scope, contadores de "
        "parámetros y variables, y el índice del cuádruplo donde inicia la función "
        "(quad_start). El scope 'global' se registra como una entrada especial."
    )
    add_code_block(doc, """class VariableTable:
    def __init__(self):
        self.vars = {}
    def add(self, name, var_type):
        if name in self.vars:
            return False
        self.vars[name] = {'type': var_type}
        return True
    def get(self, name):
        return self.vars.get(name)
    def exists(self, name):
        return name in self.vars


class FunctionDirectory:
    def __init__(self):
        self.funcs = {}
    def add(self, name, return_type, params=None):
        if name in self.funcs:
            return False
        self.funcs[name] = {
            'return_type': return_type,
            'params': params if params else [],
            'var_table': VariableTable(),
            'num_params': 0,
            'num_vars': 0,
            'quad_start': None,
        }
        return True
""")

    # SemanticManager
    doc.add_heading("Manager semántico (SemanticManager)", level=3)
    doc.add_paragraph(
        "Es la estructura central que agrupa todo el estado del análisis: el cubo, el "
        "directorio de funciones, la lista de cuádruplos, y las pilas auxiliares "
        "(operandos, tipos, saltos pendientes, breaks pendientes por ciclo, llamadas en "
        "curso, contador de argumentos). Expone helpers para emitir cuádruplos, rellenar "
        "saltos con backpatching, declarar variables y buscar identificadores con "
        "resolución por scope (local antes que global)."
    )
    add_code_block(doc, """class SemanticManager:
    def __init__(self):
        self.func_dir = FunctionDirectory()
        self.cube = SemanticCube()
        self.quads = []
        self.operand_stack = []
        self.type_stack = []
        self.jump_stack = []
        self.break_stack = []   # pila de listas: un sublist por ciclo anidado
        self.call_stack = []
        self.arg_counter = []
        self.temp_counter = 0
        self.current_scope = 'global'
        self.has_errors = False

    def new_temp(self):
        self.temp_counter += 1
        return f"t{self.temp_counter}"

    def emit(self, op, arg1, arg2, res, type_res='-'):
        self.quads.append(Quadruple(op, arg1, arg2, res, type_res))
        return len(self.quads)   # índice 1-based

    def fill(self, idx, value):
        self.quads[idx - 1].res = value

    def lookup_var(self, name):
        local = self.func_dir.get(self.current_scope)
        if local and local['var_table'].exists(name):
            return local['var_table'].get(name), self.current_scope
        glob = self.func_dir.get('global')
        if glob and glob['var_table'].exists(name):
            return glob['var_table'].get(name), 'global'
        return None, None
""")

    # Quadruple
    doc.add_heading("Cuádruplo (Quadruple)", level=3)
    doc.add_paragraph(
        "Cada cuádruplo es una tupla extendida (op, arg1, arg2, res, type_res). La columna "
        "type_res es la extensión solicitada por la entrega y contiene el tipo del "
        "resultado de la operación (int, float, string, bool) o '-' cuando la operación no "
        "produce un valor tipado (saltos, llamadas, prints)."
    )
    add_code_block(doc, """class Quadruple:
    def __init__(self, op, arg1, arg2, res, type_res=None):
        self.op = op
        self.arg1 = arg1
        self.arg2 = arg2
        self.res = res
        self.type_res = type_res if type_res is not None else '-'
""")

    doc.add_heading("2.3 Algoritmo de análisis semántico", level=2)
    doc.add_paragraph(
        "Las acciones semánticas se enganchan a las reglas de la gramática mediante el "
        "patrón clásico de neuralgic points: producciones vacías que se reducen en un punto "
        "específico de otra regla y disparan código. PLY permite acceder a los símbolos "
        "previamente reducidos en la regla padre con p[-1], p[-2], etc., lo cual hace "
        "posible inicializar el scope justo después de ver el identificador de una función "
        "o registrar un parámetro tan pronto como se ve su declaración."
    )

    doc.add_paragraph("Las cuatro validaciones semánticas se implementan así:")
    doc.add_paragraph(
        "1) Coincidencia de tipos. Cada vez que se reduce un operador binario, se "
        "consulta el cubo semántico con el operador y los tipos del tope de la "
        "pila. Si el resultado es 'error', se reporta el conflicto; en caso "
        "contrario, se genera un cuádruplo con el operador y un temporal "
        "como resultado, y se vuelve a empujar el temporal y su tipo a las pilas."
    )
    doc.add_paragraph(
        "2) Declaraciones de identificadores y scopes. Al reducir el bloque "
        "var del programa, se invoca SemanticManager.declare_var para cada "
        "identificador; si la variable ya existe en el scope actual, se "
        "reporta duplicado. Las funciones se registran al ver tipo ID en la "
        "regla de función. Cuando un identificador se usa en una expresión, "
        "se llama a lookup_var que busca primero en la tabla local del "
        "scope actual y luego en la global; si no se encuentra, se reporta "
        "no declarada."
    )
    doc.add_paragraph(
        "3) Llamadas a funciones. Al ver ID ( en una llamada, se emite un "
        "cuádruplo ERA y se crea un contador de argumentos. Cada expresión "
        "que aparece como argumento pasa por np_call_arg, que verifica el "
        "tipo contra el parámetro esperado y emite un cuádruplo PARAM. Al "
        "cerrar el paréntesis, se compara el número de argumentos recibidos "
        "contra los declarados y se emite GOSUB; si la llamada está dentro "
        "de una expresión, además se emite una asignación que copia el "
        "valor de retorno (almacenado en la variable global con el nombre "
        "de la función) a un temporal."
    )
    doc.add_paragraph(
        "4) Control de flujo. La condición de un if o while genera un "
        "valor tipo bool gracias al cubo. El neuralgic point inmediatamente "
        "después de la condición verifica que el tipo en el tope de la "
        "pila sea bool; si no, reporta error. Se emiten saltos GOTOF con "
        "destino pendiente, cuyo índice se apila en jump_stack; al cerrar "
        "el bloque se rellena el salto con la posición del próximo "
        "cuádruplo. Para break se mantiene una pila adicional break_stack "
        "que es una pila de listas, una por cada ciclo anidado; si break "
        "aparece sin un ciclo padre, se reporta error."
    )

    doc.add_heading("2.4 Representación intermedia", level=2)
    doc.add_paragraph(
        "La representación intermedia es una lista plana de cuádruplos numerados desde 1. "
        "El conjunto de instrucciones generadas se resume en la siguiente tabla:"
    )

    ir_table = doc.add_table(rows=1, cols=2)
    ir_table.style = 'Light Grid Accent 1'
    hdr = ir_table.rows[0].cells
    hdr[0].text = "Instrucción"
    hdr[1].text = "Significado"
    ir_rows = [
        ("+, -, *, /", "Operaciones aritméticas. arg1 op arg2 → res."),
        (">, <, >=, <=, ==, !=", "Operaciones relacionales. Resultado tipo bool."),
        ("=", "Asignación. arg1 → res."),
        ("PRINT", "Imprime el valor en res."),
        ("GOTO", "Salto incondicional al cuádruplo res."),
        ("GOTOF", "Salto si arg1 es falso. Destino: res."),
        ("GOTOT", "Salto si arg1 es verdadero. Usado en do-while."),
        ("ERA", "Activar registro de la función (preludio de llamada)."),
        ("PARAM", "Pasar argumento. arg1 → posición res (par1, par2, ...)."),
        ("GOSUB", "Saltar al cuádruplo de inicio de la función."),
        ("RETURN", "Marca el final de una rama con retorno."),
        ("ENDFUNC", "Marca el final de una función."),
        ("END", "Final del programa."),
    ]
    for row in ir_rows:
        cells = ir_table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    doc.add_paragraph(
        "El primer cuádruplo del programa es siempre un GOTO con destino pendiente; al "
        "ver la palabra main, el neuralgic point np_main_start rellena ese salto con la "
        "posición del primer cuádruplo del cuerpo principal. Esto garantiza que la "
        "ejecución empiece en main aunque las funciones se hayan declarado primero. Cada "
        "función registra su quad_start cuando empieza su cuerpo y termina con un "
        "ENDFUNC, lo cual permite que GOSUB salte directamente a la posición correcta."
    )


def build_casos(doc):
    doc.add_heading("3. Casos de prueba", level=1)
    doc.add_paragraph(
        "Se presentan ocho casos divididos en cuatro pares (uno por cada validación "
        "semántica), más un caso integral que combina todas las instrucciones del "
        "lenguaje. Por cada caso se muestra el código fuente en Little Duck, el "
        "resultado esperado y un extracto de la salida real del compilador."
    )

    casos = [
        ("01_tipos_exito", "Caso 1 (éxito) — Coincidencia de tipos",
         "Operaciones aritméticas y asignaciones con tipos compatibles, incluyendo "
         "conversión int → float."),
        ("02_tipos_error", "Caso 2 (fallo) — Coincidencia de tipos",
         "Asignaciones incompatibles (float a int, int a string) y operaciones entre "
         "tipos no permitidos (string + int, string * int). El compilador debe "
         "detectar todas las violaciones."),
        ("03_scopes_exito", "Caso 3 (éxito) — Declaraciones y scopes",
         "Funciones con variables locales y parámetros; uso correcto de variables "
         "globales dentro de funciones; resolución correcta de scope (local antes "
         "que global)."),
        ("04_scopes_error", "Caso 4 (fallo) — Declaraciones y scopes",
         "Variable global duplicada, parámetro duplicado, función duplicada, uso de "
         "variables no declaradas y uso de una variable local fuera de su scope."),
        ("05_llamadas_exito", "Caso 5 (éxito) — Llamadas a funciones",
         "Funciones con retorno (int, float) y funciones void. Llamadas con número y "
         "tipos de argumentos correctos. Uso del retorno dentro de expresiones."),
        ("06_llamadas_error", "Caso 6 (fallo) — Llamadas a funciones",
         "Llamada a función inexistente, número incorrecto de argumentos (de más y "
         "de menos), uso de función void en expresión, tipo de retorno incompatible "
         "y argumento con tipo incorrecto."),
        ("07_flujo_exito", "Caso 7 (éxito) — Control de flujo",
         "Ciclo do-while con if/else anidado, y uso de break con condición. "
         "Condiciones booleanas correctas en todos los puntos de decisión."),
        ("08_flujo_error", "Caso 8 (fallo) — Control de flujo",
         "Condiciones no booleanas (int, string, int+int) en if y while, y break "
         "fuera de cualquier ciclo."),
        ("09_integral", "Caso 9 (integral) — Combinación completa",
         "Programa que utiliza todas las instrucciones del lenguaje: declaración de "
         "variables globales y locales, funciones con retorno y void, llamadas "
         "anidadas, control de flujo con if/else y do-while, break condicional, "
         "operaciones aritméticas mixtas e impresión."),
    ]

    for filename, titulo, descripcion in casos:
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(descripcion)

        code, output = load_case(filename)

        doc.add_paragraph().add_run("Código Little Duck:").bold = True
        add_code_block(doc, code)

        errores = extract_errors(output)
        if errores:
            doc.add_paragraph().add_run("Errores semánticos detectados:").bold = True
            add_code_block(doc, "\n".join(errores))
        else:
            quads = extract_quads_section(output)
            if quads:
                # Acortar si es muy largo
                lines = quads.splitlines()
                if len(lines) > 32:
                    quads = "\n".join(lines[:30] + ["... (salida truncada)"])
                doc.add_paragraph().add_run(
                    "Cuádruplos generados (extracto):"
                ).bold = True
                add_code_block(doc, quads, font_size=8)

            symbols = extract_symbols_section(output)
            if symbols:
                lines = symbols.splitlines()
                if len(lines) > 24:
                    symbols = "\n".join(lines[:22] + ["... (salida truncada)"])
                doc.add_paragraph().add_run("Tabla de símbolos:").bold = True
                add_code_block(doc, symbols, font_size=8)


def build_contribuciones(doc):
    doc.add_heading("4. Desglose de contribuciones individuales", level=1)
    doc.add_paragraph(
        "Esta sección debe completarse por el equipo, indicando la participación de "
        "cada integrante en las distintas etapas de la entrega."
    )

    contrib_table = doc.add_table(rows=1, cols=2)
    contrib_table.style = 'Light Grid Accent 1'
    hdr = contrib_table.rows[0].cells
    hdr[0].text = "Integrante"
    hdr[1].text = "Contribuciones"

    for nombre in ("[Integrante 1]", "[Integrante 2]", "[Integrante 3]", "[Integrante 4]"):
        cells = contrib_table.add_row().cells
        cells[0].text = nombre
        cells[1].text = "[Describir aquí las contribuciones]"


def build_conclusiones(doc):
    doc.add_heading("5. Conclusiones", level=1)
    doc.add_paragraph(
        "La entrega 2 del compilador Little Duck integra exitosamente la fase de "
        "análisis semántico y la generación de representación intermedia sobre el "
        "trabajo previo de léxico y sintaxis. El uso de neuralgic points dentro de la "
        "gramática LALR(1) permitió implementar las acciones semánticas de manera "
        "modular sin alterar la estructura del parser. Las cuatro validaciones "
        "semánticas se verifican antes de generar los cuádruplos, lo cual evita "
        "producir representación intermedia incorrecta y permite reportar al usuario "
        "varios errores en una sola ejecución."
    )
    doc.add_paragraph(
        "Para la siguiente entrega, el siguiente paso es la generación de código "
        "objeto o la ejecución directa de los cuádruplos mediante una máquina "
        "virtual, lo cual requerirá definir las direcciones de memoria de variables "
        "y temporales, así como la mecánica de paso de parámetros por valor en "
        "tiempo de ejecución."
    )


def main():
    doc = Document()
    set_global_style(doc)

    build_portada(doc)
    build_introduccion(doc)
    build_metodologia(doc)
    build_casos(doc)
    build_contribuciones(doc)
    build_conclusiones(doc)

    output_path = ROOT / "Reporte_Entrega2.docx"
    doc.save(output_path)
    print(f"[OK] Reporte generado en {output_path}")


if __name__ == '__main__':
    main()
